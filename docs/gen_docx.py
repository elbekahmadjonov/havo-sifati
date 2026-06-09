from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

BLUE   = RGBColor(0x1e, 0x40, 0xaf)
TEAL   = RGBColor(0x0f, 0x76, 0x6e)
PURPLE = RGBColor(0x6d, 0x28, 0xd9)
GRAY   = RGBColor(0x4b, 0x55, 0x63)
DARK   = RGBColor(0x11, 0x18, 0x27)
GREEN  = RGBColor(0x05, 0x96, 0x69)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BG_BLUE   = RGBColor(0xe0, 0xf2, 0xfe)
BG_GREEN  = RGBColor(0xdc, 0xfc, 0xe7)
BG_PURPLE = RGBColor(0xf3, 0xe8, 0xff)
BG_GRAY   = RGBColor(0xf1, 0xf5, 0xf9)
BG_ORANGE = RGBColor(0xff, 0xed, 0xd5)

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hx = '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2])
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hx)
    tcPr.append(shd)

def h1(text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = color
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:color'), '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2]))
    pBdr.append(bottom)
    pPr.append(pBdr)

def h2(text, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = color

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = DARK

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK

def formula(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(1.5)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(10)
    r.font.color.rgb = PURPLE
    r.bold = True

def code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(1.0)
    for i, line in enumerate(lines):
        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(9)
        r.font.color.rgb = DARK
        if i < len(lines) - 1:
            r.add_break()

def bullet(text, color=DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(1.2)
    r = p.add_run('•  ' + text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = color

def add_table(headers, rows, hdr_color=BLUE):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, hdr_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = WHITE
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = BG_GRAY if ri % 2 == 1 else RGBColor(0xFF, 0xFF, 0xFF)
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
    doc.add_paragraph()

def info_box(text, bg=BG_BLUE, label=''):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    prefix = label + '  ' if label else ''
    r = p.add_run(prefix + text)
    r.font.size = Pt(10)
    r.font.color.rgb = DARK
    doc.add_paragraph()

# ─── MUQOVA ────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run('HAVO SIFATI MONITORINGI')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AQI Hisoblash va ML Bashorat Formulalari')
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = TEAL

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Texnik hujjat  --  Diplom ishi 2025-2026')
r.font.size = Pt(11); r.font.color.rgb = GRAY; r.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('aqi_calculator.py  *  ml_predictor.py')
r.font.name = 'Courier New'; r.font.size = Pt(10); r.font.color.rgb = PURPLE

doc.add_page_break()

# ─── 1. KIRISH ─────────────────────────────────────────────
h1('1. Kirish va Tizim Arxitekturasi', BLUE)
body('Havo sifati monitoringi tizimi ESP32 mikrokontrolleri va sensorlar yordamida havo '
     'parametrlarini olchaydi. FastAPI server olingan malumotlarni qayta ishlab, '
     'AQI (Air Quality Index) ni hisoblaydi va keyingi soatlik bashorat beradi.')
info_box('AQI -- AQSh EPA standartlariga asoslangan 0-500 oraligidagi rkaqamli korsatkich. '
         'Qanchalik yuqori bolsa, havo shuncha iflos.', BG_BLUE, 'EPA')

h2('Sensorlar va malumot oqimi')
add_table(
    ['Sensor', 'Model', 'Olmov', 'Protokol', "AQI ta'siri"],
    [
        ['Chang sensori',    'PMS5003', 'PM2.5, PM10 (mcg/m3)', 'UART2',         'Asosiy (ustun)'],
        ['Gaz sensori 1',    'MQ-135',  'CO2, NH3, Benzol (DO)', 'Digital',       'Zaxira'],
        ['Gaz sensori 2',    'MQ-2',    'Metan, LPG, Tutun (DO)', 'Digital',      'Zaxira'],
        ['Gaz sensori 3',    'MQ-7',    'Uglerod oksidi CO (DO)', 'Digital',       'Zaxira'],
        ['Harorat/Namlik',   'DHT22',   'C, %',                  'DHT protokol',  'Bonus'],
        ['Bosim sensori',    'BMP280',  'hPa',                   'I2C 0x76',      'Korreksiya'],
    ],
    BLUE
)

doc.add_page_break()

# ─── 2. EPA FORMULA ────────────────────────────────────────
h1('2. EPA Interpolatsiya Formulasi', BLUE)
body('Barcha PM sensorlar uchun AQI qiymati AQSh EPA standartidagi linear interpolatsiya '
     'formulasi orqali hisoblanadi:')

formula('         (AQI_high - AQI_low)')
formula('AQI  =  ========================  x  (C - C_low)  +  AQI_low')
formula('         (C_high  - C_low)   ')

h3("O'zgaruvchilar:")
bullet('C        -- Olchangan haqiqiy konsentratsiya (mcg/m3)')
bullet('C_low    -- C tushgan oraliqning pastki chegarasi')
bullet('C_high   -- C tushgan oraliqning yuqori chegarasi')
bullet('AQI_low  -- Shu oraliqning AQI pastki qiymati')
bullet('AQI_high -- Shu oraliqning AQI yuqori qiymati')

h3('Python kodi:')
code_block([
    'def _epa_interpolatsiya(c: float, breakpoints: list) -> Optional[int]:',
    '    for c_min, c_max, i_min, i_max in breakpoints:',
    '        if c_min <= c <= c_max:',
    '            return round((i_max-i_min)/(c_max-c_min)*(c-c_min)+i_min)',
    '    return 500 if c > breakpoints[-1][1] else None',
])

info_box("Agar konsentratsiya barcha oraliqdan oshib ketsa (C > 500.4), "
         "funksiya 500 qaytaradi -- bu maksimal xavfli daraja.", BG_ORANGE, 'Chegaraviy holat')

# ─── 3. PM2.5 ──────────────────────────────────────────────
h1('3. PM2.5 --> AQI Hisoblash (EPA Breakpoints)', TEAL)
body('PM2.5 -- diametri 2.5 mkm dan kichik mayda zarrachalar. '
     "O'pka ichiga kirib, jiddiy kasalliklarga olib kelishi mumkin.")

add_table(
    ['C_low (mcg/m3)', 'C_high', 'AQI_low', 'AQI_high', 'Daraja'],
    [
        ['0.0',   '12.0',  '0',   '50',  'Yaxshi'],
        ['12.1',  '35.4',  '51',  '100', "O'rtacha"],
        ['35.5',  '55.4',  '101', '150', 'Sezgir guruh uchun zararli'],
        ['55.5',  '150.4', '151', '200', 'Zararli'],
        ['150.5', '250.4', '201', '300', 'Juda zararli'],
        ['250.5', '350.4', '301', '400', '--'],
        ['350.5', '500.4', '401', '500', 'Xavfli'],
    ],
    TEAL
)

h3('Hisoblash misoli -- PM2.5 = 40 mcg/m3:')
formula('Oraliq: 35.5 <= 40 <= 55.4  -->  AQI oraliq: 101-150')
formula('')
formula('         (150 - 101)')
formula('AQI  =  =============  x  (40 - 35.5)  +  101')
formula('         (55.4 - 35.5)')
formula('')
formula('      =  (49 / 19.9)  x  4.5  +  101')
formula('      =  2.462  x  4.5  +  101  =  112  (Sezgir guruh)')

h2('PM10 Breakpoints jadvali')
add_table(
    ['C_low (mcg/m3)', 'C_high', 'AQI_low', 'AQI_high', 'Daraja'],
    [
        ['0',   '54',  '0',   '50',  'Yaxshi'],
        ['55',  '154', '51',  '100', "O'rtacha"],
        ['155', '254', '101', '150', 'Sezgir guruh'],
        ['255', '354', '151', '200', 'Zararli'],
        ['355', '424', '201', '300', 'Juda zararli'],
        ['425', '504', '301', '400', '--'],
        ['505', '604', '401', '500', 'Xavfli'],
    ],
    TEAL
)

doc.add_page_break()

# ─── 4. MQ ─────────────────────────────────────────────────
h1('4. MQ Sensorlar --> AQI (Ball Tizimi)', PURPLE)
body("MQ sensorlar faqat raqamli (DO) signal beradi: 1=toza, 0=gaz aniqlandi. "
     "PM sensor bo'lmasa zaxira sifatida ishlatiladi.")

h2("Asosiy formula -- Iflos sensorlar soni --> AQI")
formula('iflos_soni = count(mq135==0) + count(mq2==0) + count(mq7==0)')
formula('')
formula('  { 0 iflos  -->  AQI = 30  }   (Yaxshi)')
formula('  { 1 iflos  -->  AQI = 75  }   (O\'rtacha)')
formula('  { 2 iflos  -->  AQI = 125 }   (Sezgir guruh)')
formula('  { 3 iflos  -->  AQI = 175 }   (Zararli)')

add_table(
    ['Iflos sensorlar', 'AQI', 'Daraja', 'Izoh'],
    [
        ['0 ta (hammasi toza)', '30',  'Yaxshi',       'Barcha DO=1'],
        ['1 ta',                '75',  "O'rtacha",     'Bitta DO=0'],
        ['2 ta',                '125', 'Sezgir guruh', 'Ikkita DO=0'],
        ['3 ta',                '175', 'Zararli',      'Uchala DO=0'],
    ],
    PURPLE
)

h3('Python kodi:')
code_block([
    '_MQ_AQI_JADVAL = { 0: 30, 1: 75, 2: 125, 3: 175 }',
    '',
    'def _mq_aqi(mq135, mq2, mq7, harorat, namlik):',
    '    iflos = sum(1 for val in [mq135, mq2, mq7]',
    '                if val is not None and val == 0)',
    '    aqi = _MQ_AQI_JADVAL.get(iflos, 175)',
    '    if iflos > 0:',
    '        if harorat is not None and harorat > 35.0:',
    '            aqi += 25   # harorat bonusi',
    '        if namlik is not None and namlik > 80.0:',
    '            aqi += 15   # namlik bonusi',
    '    return int(min(aqi, 500))',
])

add_table(
    ['Sensor', 'GPIO', 'Aniqlaydigan gazlar', "DO=0 ma'nosi"],
    [
        ['MQ-135', 'GPIO 4',  'CO2, NH3, Benzol, Spirt',     'Zararli gaz oshdi'],
        ['MQ-2',   'GPIO 5',  'Metan, LPG, Tutun, Propan',   'Yonuvchan gaz bor'],
        ['MQ-7',   'GPIO 19', 'Uglerod oksidi (CO)',          'CO bor -- ENG XAVFLI!'],
    ],
    PURPLE
)

doc.add_page_break()

# ─── 5. IQLIM ──────────────────────────────────────────────
h1("5. Iqlim Bonuslari va Bosim Korreksiyasi", TEAL)

h2("5.1 Harorat va Namlik Bonuslari (MQ rejimida)")
body("Gaz aniqlanganda (iflos>0) yuqori harorat va namlik gazning ta'sirini kuchaytiradi:")

formula('if (iflos_soni > 0) AND (harorat > 35 C):  AQI_mq  +=  25')
formula('if (iflos_soni > 0) AND (namlik  > 80 %):  AQI_mq  +=  15')
formula('')
formula('Ikkala shart bir vaqtda:  AQI_mq  +=  25 + 15  =  +40')

add_table(
    ['Shart', 'Chegara', 'Bonus', 'Sababi'],
    [
        ['Harorat (gaz bor)', '> 35 C',  '+25 AQI', "Issiqda gazlar tez ta'sir qiladi"],
        ['Namlik (gaz bor)',  '> 80 %',  '+15 AQI', 'Yuqori namlik nafas yollarini zaiflashtiradi'],
        ['Ikkala shart',      'ikkalasi', '+40 AQI', "Kombinatsion ta'sir"],
    ],
    TEAL
)
info_box("Bonuslar FAQAT MQ rejimida (PM yo'q bo'lganda) va gaz aniqlangandagina ishlaydi.",
         BG_BLUE, 'Shart')

h2("5.2 Bosim Korreksiyasi (har doim ishlaydi)")
formula('if bosim < 990 hPa:   AQI_yakuniy  +=  +10')
formula('if bosim > 1030 hPa:  AQI_yakuniy  -=   -5')
formula('else (990-1030):      AQI_yakuniy  +=    0   (ozgarishsiz)')

add_table(
    ['Bosim (hPa)', "O'zgarish", 'Sababi'],
    [
        ['< 990',    '+10 AQI', "Past bosimda gaz tarqalmaydi, to'planib qoladi"],
        ['990-1030', '0',       'Normal atmosfera bosimi'],
        ['> 1030',   '-5 AQI',  'Yuqori bosimda havo yaxshi aylanadi'],
    ],
    TEAL
)

doc.add_page_break()

# ─── 6. YAKUNIY AQI ────────────────────────────────────────
h1("6. Yakuniy AQI -- Ustunlik Zanjiri", BLUE)

code_block([
    '                  PM2.5 mavjudmi?  PM10 mavjudmi?',
    '                          |',
    '            ┌─────────────┴─────────────┐',
    '           HA                          YOQ',
    '            |                           |',
    '  AQI = max(pm25_aqi,         MQ ball tizimi',
    '            pm10_aqi)        + iqlim bonuslari',
    '    (EPA interpolatsiya)          |',
    '            |                    |',
    '            └──────────┬─────────┘',
    '                       |',
    '            + Bosim korreksiyasi',
    '                       |',
    '            AQI = clamp(0 ... 500)',
])

h3('Python kodi (hisobla_aqi):')
code_block([
    'def hisobla_aqi(mq135, mq2, mq7, harorat, namlik, bosim, pm25, pm10):',
    '    pm_aqilar = []',
    '    if pm25 is not None and pm25 >= 0:',
    '        v = _pm25_aqi(pm25)',
    '        if v is not None: pm_aqilar.append(v)',
    '    if pm10 is not None and pm10 >= 0:',
    '        v = _pm10_aqi(pm10)',
    '        if v is not None: pm_aqilar.append(v)',
    '',
    '    if pm_aqilar:',
    '        aqi = max(pm_aqilar)   # PM ustun -- MQ etiborga olinmaydi',
    '    else:',
    '        aqi = _mq_aqi(mq135, mq2, mq7, harorat, namlik)',
    '',
    '    if bosim is not None:',
    '        if bosim < 990.0:    aqi += 10',
    '        elif bosim > 1030.0: aqi -= 5',
    '',
    '    return int(min(max(aqi, 0), 500))',
])

h2('Hisoblash misollari:')
add_table(
    ['Holat', 'Kiritma', 'Hisoblash', 'AQI'],
    [
        ['Faqat PM2.5',         'pm25=40',               'EPA: (49/19.9)x4.5+101',   '112'],
        ['PM2.5 past',          'pm25=5',                'EPA: (50/12)x5',             '21'],
        ['Hamma toza, PM yoq',  'mq x=1, PM yoq',        'MQ: 0 iflos -> 30',          '30'],
        ['1 sensor gaz',        'mq135=0',               'MQ: 1 iflos -> 75',          '75'],
        ['2 sensor + issiq',    'mq135=0, mq2=0, t=37',  '2 iflos=125, +25 bonus',    '150'],
        ['PM bor + gaz bor',    'pm25=20, mq135=0',      'max(79) -- MQ ignored',      '79'],
        ['PM + past bosim',     'pm25=40, bosim=985',    '112 + 10 bosim',            '122'],
        ['Hamma iflos+issiq+nam','mq x=0, t=38, n=85',   '3 iflos=175, +25+15',       '215'],
    ],
    BLUE
)

doc.add_page_break()

# ─── 7. AQI DARAJALARI ─────────────────────────────────────
h1("7. AQI Darajalari va Sog'liq Tavsiyalari", GREEN)

add_table(
    ['AQI', 'Daraja', "Sog'liq ta'siri", 'Asosiy tavsiya'],
    [
        ['0-50',    'Yaxshi',                  'Hech qanday xavf yoq',         'Erkin faoliyat'],
        ['51-100',  "O'rtacha",                "Sezgirlarga yengil ta'sir",    'Mashqni cheklang'],
        ['101-150', 'Sezgir guruh zararli',    "Bolalar, keksalar ta'sirlanadi",'Sezgirlar chiqmasin'],
        ['151-200', 'Zararli',                 "Barcha ta'sirlanadi",          'Tashqarini kamaytir'],
        ['201-300', 'Juda zararli',            'Jiddiy xavf',                  'Niqob majburiy'],
        ['301-500', 'Xavfli',                  'Favqulodda holat',             'Tibbiy yordam'],
    ],
    GREEN
)

doc.add_page_break()

# ─── 8. ML PREDICTOR ───────────────────────────────────────
h1('8. ML Bashorat Modeli (ml_predictor.py)', PURPLE)
body('Tizim ikkita bashorat rejimini qollabquvvatlaydi: LSTM neyron tarmogi (asosiy) '
     'va statistik zaxira (LSTM model topilmaganda avtomatik yoqiladi).')

h2("8.1 LSTM Neyron Tarmog'i")
body('LSTM (Long Short-Term Memory) -- vaqt qatorlari uchun moljallangan rekurrent neyron tarmogi. '
     "Songi N ta olchovga asoslanib keyingi 1 soatlik AQI ni bashorat qiladi.")

h3('Arxitektura:')
code_block([
    'Kirish:   [AQI_t-N, AQI_t-N+1, ..., AQI_t]  <--  N ta ketma-ket olmov',
    'Qayta     LSTM qatlam 1 --> LSTM qatlam 2 --> Dense',
    'ishlash:  (vaqt bogliqlini eslaydi)',
    'Chiqish:  AQI_t+1  <--  keyingi 1 soat bashorati',
])
bullet("Kirish xususiyatlar: aqi, pm25, pm10, harorat, namlik, bosim")
bullet("O'qitish: python ml/train_model.py")
bullet("Model fayli: ml/models/lstm_model.keras")

h2('8.2 Statistik Zaxira Bashorati')
body("LSTM model topilmasa yoki malumot yetarli bolmasa, og'irlikli ortacha formula ishlatiladi:")

formula("w_i  =  1.2^i     (i = 0,1,...,N-1  eski->yangi)")
formula('')
formula("         sum(AQI_i * w_i)")
formula("AQI* =  =================")
formula("            sum(w_i)     ")
formula('')
formula("Ishonch  =  max(0.30,  min(0.75,  1.0 - std(AQI) / 120))")

h3('Python kodi:')
code_block([
    "og_irliklar = [1.2 ** i for i in range(len(aqilar))]",
    "",
    "bashorat = round(",
    "    sum(a * w for a, w in zip(aqilar, og_irliklar))",
    "    / sum(og_irliklar)",
    ")",
    "",
    "std     = statistics.stdev(aqilar)",
    "ishonch = max(0.30, min(0.75, 1.0 - std / 120))",
])
bullet("Geometrik og'irlik (1.2^i): yangi o'lchovlar 20% ko'proq ta'sir qiladi")
bullet("std/120: AQI tebranishi katta bolsa ishonch kamayadi")
bullet("Ishonch 0.30-0.75 oraligida cheklanadi")

h2('8.3 Anomaliya Aniqlash (Z-ball)')
formula("         | AQI_hozir - mean(AQI_tarix) |")
formula("z  =  =====================================")
formula("              std(AQI_tarix)             ")
formula('')
formula("z > 3  -->  Anomaliya (jiddiy chekinish)")
formula("z > 2  -->  O'rta daraja")
formula("z <= 2 -->  Normal")

h3('Python kodi:')
code_block([
    "orta = statistics.mean(aqilar)",
    "std  = statistics.stdev(aqilar)",
    "z    = abs(hozirgi['aqi'] - orta) / max(std, 0.001)",
    "",
    "if z > 3: daraja = 'yuqori'   # anomaliya",
    "if z > 2: daraja = 'o\\'rta'",
    "else:     daraja = 'normal'",
])

h2('8.4 Bashorat holatlari va ishonch darajalari')
add_table(
    ['Holat', 'Usul', 'Ishonch', 'Xabar'],
    [
        ["Ma'lumot yo'q",    'yetarli_malumot_yoq', '0%',   "Sensor malumotlari kutilmoqda"],
        ["1 ta o'lchov",     'kam_malumot',          '0%',   'Kamida 2 ta kerak'],
        ["2-11 ta o'lchov",  'statistik_ogirlikli',  '30-75%', "So'nggi N ta olchovga asoslangan"],
        [">= 5 + LSTM bor",  'lstm_model',           '75%+', "LSTM neyron tarmog'i bashorati"],
    ],
    PURPLE
)

h2('8.5 Trend aniqlash')
formula("farq = AQI_bashorat - AQI_hozir")
formula('')
formula("farq > +10  -->  Yomonlashadi  (qizil)")
formula("farq < -10  -->  Yaxshilanadi  (yashil)")
formula("-10 <= farq <= +10  -->  Barqaror  (kulrang)")

doc.add_page_break()

# ─── 9. XULOSA ─────────────────────────────────────────────
h1('9. Xulosa -- Tolik Formula Zanjiri', BLUE)

code_block([
    "ESP32 sensorlar olchaydi",
    "    |",
    "POST /api/sensor  (JSON)",
    "    |",
    "hisobla_aqi(pm25, pm10, mq135, mq2, mq7, harorat, namlik, bosim)",
    "    |-- PM2.5/PM10 bor?  -->  max(_pm25_aqi, _pm10_aqi)  [EPA]",
    "    +-- PM yoq?          -->  _mq_aqi + iqlim bonuslari  [Ball]",
    "                                       |",
    "                          + bosim korreksiyasi",
    "                                       |",
    "                          AQI = clamp(0..500)",
    "    |",
    "SQLite ga saqlash  (havo_data.db)",
    "    |",
    "_statistik_bashorat()  yoki  LSTM.predict()",
    "    |",
    "Dashboard + Telegram ogohlantirish",
])

info_box(
    "PM sensorlar AQI ni EPA standartiga kora aniq hisoblaydi. "
    "MQ sensorlar PM bolmasa zaxira sifatida ishlaydi. "
    "LSTM model vaqt qatoriga asoslanib keyingi soat AQI ni bashorat qiladi.",
    BG_GREEN, 'Xulosa'
)

OUT = r'c:\Users\seed\Documents\diplom_server\docs\AQI_Formulalar.docx'
doc.save(OUT)
print('SAQLANDI:', OUT)
